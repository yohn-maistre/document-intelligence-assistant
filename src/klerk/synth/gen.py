"""Synthetic corpus generator — Nemotron-grounded JSON → PDF/DOCX/MD.

Pipeline per doc:
  1. Build a system+user prompt from the DocSpec (klerk.synth.specs.CORPUS).
  2. Call Nemotron via the cached router (DiskCache hits make regen free).
  3. Parse the JSON response into a `DocBody` (title + sections + optional table).
  4. Render via the format writer: PDF (reportlab), DOCX (python-docx), or MD.
  5. Write to `out_dir/{doc_id}.{ext}` and append a row to manifest.json.

The JSON-output prompt is deliberately strict so the parse step is robust:
the model returns `{"title", "sections": [{"heading", "paragraphs": [str]}],
"table": {"headers", "rows"} | null}`. The generator validates with Pydantic
and falls back to a clean error rather than a corrupted doc.
"""

from __future__ import annotations

import json
import os
from dataclasses import asdict
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field, ValidationError

from klerk.synth.specs import CORPUS, DocSpec


# ─── Output schema (the model is asked to populate this exactly) ─────────────
class DocSection(BaseModel):
    heading: str
    paragraphs: list[str] = Field(default_factory=list)


class DocTable(BaseModel):
    headers: list[str]
    rows: list[list[str]]


class DocBody(BaseModel):
    title: str
    sections: list[DocSection]
    table: DocTable | None = None


# ─── Prompt builders ─────────────────────────────────────────────────────────
SYSTEM_TEMPLATE = """\
You write realistic fictional corporate documents for PT Fata Organa Solusi,
a fictional Indonesian-Japanese technology consulting firm based in Jakarta
with a satellite office in Tokyo. CAC Holding Japan is one of their largest
clients.

You will be given a document spec. Produce the document body as STRICT JSON
that matches this schema:

  {
    "title": "<the document title>",
    "sections": [
      {
        "heading": "<section heading>",
        "paragraphs": ["<paragraph 1>", "<paragraph 2>", ...]
      },
      ...
    ],
    "table": {"headers": [...], "rows": [[...], ...]} | null
  }

Rules:
  - Output ONLY the JSON object. No prose preamble, no markdown fences.
  - Use 3-6 sections, each with 1-3 paragraphs.
  - Paragraphs are plain text (no markdown). Keep them tight and realistic.
  - If the spec mentions a table, include the `table` field. Otherwise null.
  - Include realistic placeholder names (Yan, Putri, Galih for Indonesian
    staff; Tanaka, Yamada, Sato for Japanese staff).
  - Use realistic currency placeholders (IDR ..., JPY ...) and dates.
  - If the spec mentions a contradiction with another document, make the
    contradicting claim concrete and unambiguous so a contradiction
    scanner can flag it (don't bury the conflict in vague language).
  - If the spec mentions cross-references, weave them in naturally:
    "see hr_consultant_rate_card_2025 for the rate table".
"""


def _user_prompt(spec: DocSpec) -> str:
    bits = [
        f"DOC ID: {spec.doc_id}",
        f"CATEGORY: {spec.category}",
        f"TITLE: {spec.title}",
        f"LOCALE: {spec.locale}",
        f"FORMAT TARGET: {spec.format}",
    ]
    if spec.date_stamp:
        bits.append(f"DATE STAMP: {spec.date_stamp}")
    if spec.has_table:
        bits.append("REQUIREMENT: include a structured table.")
    if spec.contradiction_pair:
        bits.append(
            f"CONTRADICTION: this document is part of a deliberately "
            f"contradicting pair with {spec.contradiction_pair[1]}. Make "
            f"the conflicting claim concrete."
        )
    if spec.cross_refs:
        bits.append(
            "CROSS-REFERENCES TO WEAVE IN: " + ", ".join(spec.cross_refs)
        )
    bits.append("")
    bits.append("SPEC:")
    bits.append(spec.brief)
    bits.append("")
    bits.append("Return the JSON body now.")
    return "\n".join(bits)


# ─── LLM call ────────────────────────────────────────────────────────────────
def _call_llm(spec: DocSpec) -> DocBody:
    """Run Nemotron with cache; parse + validate the JSON response."""
    from klerk.llm.router import complete

    messages = [
        {"role": "system", "content": SYSTEM_TEMPLATE},
        {"role": "user", "content": _user_prompt(spec)},
    ]
    response = complete(
        messages=messages,
        locale=spec.locale,
        temperature=0.4,
        max_tokens=2200,
        response_format={"type": "json_object"},
    )
    text = response.choices[0].message.content or ""
    return _parse_body(text, spec)


def _parse_body(text: str, spec: DocSpec) -> DocBody:
    """Robust JSON → DocBody parse. Strips code fences if the model added them."""
    cleaned = text.strip()
    if cleaned.startswith("```"):
        # Strip the first line (```json or ```) and the trailing fence.
        cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned
        cleaned = cleaned.rsplit("```", 1)[0].strip()
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as e:
        raise RuntimeError(
            f"{spec.doc_id}: LLM returned non-JSON body. First 200 chars: "
            f"{cleaned[:200]!r}. Parse error: {e}"
        ) from e
    try:
        return DocBody.model_validate(data)
    except ValidationError as e:
        raise RuntimeError(
            f"{spec.doc_id}: JSON didn't match DocBody schema: {e}"
        ) from e


# ─── Format writers ──────────────────────────────────────────────────────────
def _write_md(path: Path, body: DocBody, spec: DocSpec) -> None:
    lines: list[str] = [f"# {body.title}", ""]
    if spec.date_stamp:
        lines.append(f"_{spec.date_stamp}_")
        lines.append("")
    for sec in body.sections:
        lines.append(f"## {sec.heading}")
        lines.append("")
        for p in sec.paragraphs:
            lines.append(p)
            lines.append("")
    if body.table:
        lines.append("## Table")
        lines.append("")
        lines.append("| " + " | ".join(body.table.headers) + " |")
        lines.append("|" + "|".join(["---"] * len(body.table.headers)) + "|")
        for row in body.table.rows:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_docx(path: Path, body: DocBody, spec: DocSpec) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(body.title, level=0)
    if spec.date_stamp:
        para = doc.add_paragraph()
        run = para.add_run(spec.date_stamp)
        run.italic = True
    for sec in body.sections:
        doc.add_heading(sec.heading, level=1)
        for p in sec.paragraphs:
            doc.add_paragraph(p)
    if body.table:
        doc.add_heading("Table", level=1)
        table = doc.add_table(rows=1 + len(body.table.rows), cols=len(body.table.headers))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(body.table.headers):
            hdr[i].text = h
        for r_idx, row in enumerate(body.table.rows, start=1):
            cells = table.rows[r_idx].cells
            for c_idx, val in enumerate(row):
                if c_idx < len(cells):
                    cells[c_idx].text = val
    doc.save(str(path))


def _write_pdf(path: Path, body: DocBody, spec: DocSpec) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    flowables: list[Any] = [Paragraph(body.title, styles["Title"])]
    if spec.date_stamp:
        flowables.append(Paragraph(f"<i>{spec.date_stamp}</i>", styles["Italic"]))
    flowables.append(Spacer(1, 12))

    for sec in body.sections:
        flowables.append(Paragraph(sec.heading, styles["Heading1"]))
        for p in sec.paragraphs:
            flowables.append(Paragraph(p, styles["BodyText"]))
        flowables.append(Spacer(1, 8))

    if body.table:
        flowables.append(Paragraph("Table", styles["Heading1"]))
        data = [body.table.headers, *body.table.rows]
        table = Table(data, repeatRows=1)
        table.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
            ])
        )
        flowables.append(table)

    pdf = SimpleDocTemplate(str(path), pagesize=LETTER, title=body.title)
    pdf.build(flowables)


_WRITERS = {
    "md": _write_md,
    "docx": _write_docx,
    "pdf": _write_pdf,
}


# ─── Orchestrator ────────────────────────────────────────────────────────────
def generate_one(spec: DocSpec, out_dir: Path) -> Path:
    """Generate + write one doc. Returns the written path."""
    body = _call_llm(spec)
    ext = spec.format
    path = out_dir / f"{spec.doc_id}.{ext}"
    path.parent.mkdir(parents=True, exist_ok=True)
    _WRITERS[ext](path, body, spec)
    return path


def generate_corpus(
    out_dir: Path,
    *,
    corpus: list[DocSpec] | None = None,
    skip_existing: bool = True,
) -> list[dict[str, Any]]:
    """Generate every doc in the corpus plan. Returns manifest rows.

    Cache-friendly: skip_existing=True (default) means previously generated
    files survive across runs; pair with the LLM router's DiskCache for
    cheap regen. Set to False to force regenerate everything.
    """
    corpus = corpus or CORPUS
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, Any]] = []
    errors: list[str] = []
    for spec in corpus:
        target = out_dir / f"{spec.doc_id}.{spec.format}"
        if skip_existing and target.exists():
            manifest_rows.append({**asdict(spec), "path": str(target), "status": "cached"})
            continue
        try:
            written = generate_one(spec, out_dir)
            manifest_rows.append({**asdict(spec), "path": str(written), "status": "generated"})
        except Exception as e:  # noqa: BLE001
            errors.append(f"{spec.doc_id}: {type(e).__name__}: {e}")
            manifest_rows.append({**asdict(spec), "path": None, "status": "failed", "error": str(e)})

    manifest_path = out_dir / "manifest.json"
    manifest_payload = {
        "corpus_root": str(out_dir),
        "n_specs": len(corpus),
        "n_generated": sum(1 for r in manifest_rows if r["status"] == "generated"),
        "n_cached": sum(1 for r in manifest_rows if r["status"] == "cached"),
        "n_failed": sum(1 for r in manifest_rows if r["status"] == "failed"),
        "docs": manifest_rows,
        "errors": errors,
    }
    manifest_path.write_text(json.dumps(manifest_payload, indent=2, ensure_ascii=False))
    return manifest_rows
