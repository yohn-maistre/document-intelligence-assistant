"""Unified parser — one entry point, format-aware dispatch.

The `ParsedDocument` dataclass is the shape that flows into the chunker.
Markdown is preserved as-is; PDF/DOCX get Docling's layout-aware extraction
into Markdown so downstream chunking and embedding sees consistent text.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


_DOCLING_FORMATS = {".pdf", ".docx", ".pptx", ".xlsx", ".html", ".htm"}
_NATIVE_TEXT = {".md", ".markdown", ".txt"}


@dataclass
class ParsedDocument:
    """Output of `parse()`. `text` is the canonical content for downstream use."""

    source: Path
    text: str
    locale: str = "und"  # ISO 639-1: en, id, und (undetermined)
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def doc_id(self) -> str:
        """Stable identifier for the doc (file stem). Caller may override in meta."""
        return self.meta.get("doc_id", self.source.stem)


def parse(path: str | Path) -> ParsedDocument:
    """Parse one file. Raises FileNotFoundError if missing.

    Routing is by extension; override the parser with KLERK_PARSER=pymupdf
    if Docling's heavy extras fail to install.
    """
    p = Path(path).expanduser().resolve()
    if not p.exists():
        raise FileNotFoundError(f"parse: file not found: {p}")

    suffix = p.suffix.lower()
    parser = os.environ.get("KLERK_PARSER", "docling")

    if suffix in _NATIVE_TEXT:
        return _parse_native(p, suffix)

    # PDF: PyMuPDF when requested or as a Docling fallback (lite installs).
    if suffix == ".pdf":
        if parser == "pymupdf":
            return _parse_pymupdf(p)
        try:
            return _parse_docling(p)
        except Exception:  # noqa: BLE001
            return _parse_pymupdf(p)

    # DOCX: native python-docx unless Docling is explicitly chosen AND present.
    # python-docx reads Unicode from the XML → CJK/Bahasa extract cleanly,
    # font/version-independent (PDF glyph extraction is not).
    if suffix == ".docx":
        if parser != "docling":
            return _parse_docx(p)
        try:
            return _parse_docling(p)
        except Exception:  # noqa: BLE001
            return _parse_docx(p)

    if suffix in _DOCLING_FORMATS:
        try:
            return _parse_docling(p)
        except Exception:  # noqa: BLE001
            return _parse_native(p, suffix)

    # Unknown extension — try Docling (it handles many formats), fall back to text read
    try:
        return _parse_docling(p)
    except Exception:  # noqa: BLE001
        return _parse_native(p, suffix)


# ─── Native text reader ──────────────────────────────────────────────────────
def _parse_native(p: Path, suffix: str) -> ParsedDocument:
    text = p.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(
        source=p,
        text=text,
        locale=_sniff_locale(text),
        meta={"parser": "native", "suffix": suffix, "bytes": p.stat().st_size},
    )


# ─── Docling ──────────────────────────────────────────────────────────────────
def _parse_docling(p: Path) -> ParsedDocument:
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(str(p))
    markdown = result.document.export_to_markdown()
    return ParsedDocument(
        source=p,
        text=markdown,
        locale=_sniff_locale(markdown),
        meta={
            "parser": "docling",
            "suffix": p.suffix.lower(),
            "page_count": getattr(result.document, "num_pages", None),
        },
    )


# ─── PyMuPDF fallback ─────────────────────────────────────────────────────────
def _parse_pymupdf(p: Path) -> ParsedDocument:
    import pymupdf  # type: ignore[import-not-found]

    doc = pymupdf.open(p)
    pages = [page.get_text() for page in doc]
    text = "\n\n".join(pages)
    doc.close()
    return ParsedDocument(
        source=p,
        text=text,
        locale=_sniff_locale(text),
        meta={"parser": "pymupdf", "suffix": ".pdf", "page_count": len(pages)},
    )


# ─── Native DOCX reader (python-docx) ─────────────────────────────────────────
def _parse_docx(p: Path) -> ParsedDocument:
    """DOCX via python-docx — paragraphs + tables, no torch/Docling.

    Reads the document's Unicode text straight from the XML, so CJK / Bahasa
    text extracts cleanly regardless of fonts or library versions.
    """
    from docx import Document

    d = Document(str(p))
    parts: list[str] = [para.text for para in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    text = "\n".join(s for s in parts if s and s.strip())
    return ParsedDocument(
        source=p,
        text=text,
        locale=_sniff_locale(text),
        meta={"parser": "python-docx", "suffix": ".docx"},
    )


# ─── Locale sniff (cheap heuristic) ──────────────────────────────────────────
_ID_MARKERS = (
    "yang ", "dengan ", "untuk ", "tidak ", "adalah ", "akan ",
    "ini ", "itu ", "pada ", "dari ", "dalam ", "Pasal ",
    "kontrak", "perusahaan", "ketentuan", "kebijakan", "pihak",
)
_EN_MARKERS = (
    " the ", " and ", " for ", " with ", " this ", " that ",
    " from ", " policy ", " parental ", " consultant ", " employee ",
)


def _sniff_locale(text: str, sample: int = 4096) -> str:
    """Cheap markers-based locale sniff. Returns 'en', 'id', or 'und'.

    Not for production accuracy — just enough to set a default for the
    Bahasa eval split and the --locale routing.
    """
    snippet = text[:sample].lower()
    id_hits = sum(snippet.count(m) for m in _ID_MARKERS)
    en_hits = sum(snippet.count(m) for m in _EN_MARKERS)
    if id_hits > en_hits * 1.2:
        return "id"
    if en_hits > id_hits * 1.2:
        return "en"
    if id_hits + en_hits > 0:
        return "und"  # bilingual / undetermined
    return "und"
