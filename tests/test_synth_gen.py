"""Synthetic-corpus generator tests.

LLM calls are mocked end to end via a fake `complete()` returning a
canonical DocBody JSON. Format writers (PDF / DOCX / MD) are exercised
against tmp_path and verified for basic structural validity.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from klerk.synth import gen as gen_mod
from klerk.synth.gen import (
    DocBody,
    DocSection,
    DocTable,
    _parse_body,
    _user_prompt,
    _write_docx,
    _write_md,
    _write_pdf,
    generate_corpus,
)
from klerk.synth.specs import CORPUS, DocSpec, constraint_check


# ─── Corpus plan satisfies the brief ─────────────────────────────────────────
def test_corpus_satisfies_all_brief_constraints():
    results = constraint_check()
    failed = {k: v for k, v in results.items() if isinstance(v, bool) and not v}
    assert not failed, f"Constraints not met: {failed}"


def test_corpus_total_is_in_range():
    assert 25 <= len(CORPUS) <= 30


def test_corpus_doc_ids_are_unique():
    ids = [d.doc_id for d in CORPUS]
    assert len(ids) == len(set(ids))


def test_contradiction_pairs_are_symmetric():
    """If A points at B as a contradiction pair, B should point back at A."""
    pair_map = {
        d.doc_id: d.contradiction_pair[1]
        for d in CORPUS
        if d.contradiction_pair
    }
    for a, b in pair_map.items():
        assert b in pair_map, f"{a} → {b} but {b} has no contradiction_pair"
        assert pair_map[b] == a, f"Asymmetric pair: {a} ↔ {b}"


def test_cross_refs_point_to_real_docs():
    ids = {d.doc_id for d in CORPUS}
    for d in CORPUS:
        for ref in d.cross_refs:
            assert ref in ids, f"{d.doc_id} references unknown doc: {ref}"


# ─── Prompt assembly ─────────────────────────────────────────────────────────
def _make_spec(**overrides) -> DocSpec:
    base = {
        "doc_id": "test_doc",
        "category": "hr",
        "format": "md",
        "locale": "en",
        "title": "Test Doc",
        "brief": "A short test brief.",
    }
    base.update(overrides)
    return DocSpec(**base)


def test_user_prompt_includes_basic_fields():
    spec = _make_spec()
    prompt = _user_prompt(spec)
    assert "test_doc" in prompt
    assert "hr" in prompt
    assert "Test Doc" in prompt
    assert "A short test brief." in prompt


def test_user_prompt_calls_out_table_requirement():
    prompt = _user_prompt(_make_spec(has_table=True))
    assert "structured table" in prompt


def test_user_prompt_calls_out_contradiction():
    spec = _make_spec(contradiction_pair=("test_doc", "other_doc"))
    prompt = _user_prompt(spec)
    assert "CONTRADICTION" in prompt
    assert "other_doc" in prompt


def test_user_prompt_lists_cross_refs():
    prompt = _user_prompt(_make_spec(cross_refs=["doc_a", "doc_b"]))
    assert "CROSS-REFERENCES" in prompt
    assert "doc_a" in prompt
    assert "doc_b" in prompt


# ─── JSON parsing ────────────────────────────────────────────────────────────
def test_parse_body_accepts_plain_json():
    payload = {
        "title": "Hi",
        "sections": [{"heading": "S1", "paragraphs": ["p1"]}],
        "table": None,
    }
    body = _parse_body(json.dumps(payload), _make_spec())
    assert body.title == "Hi"
    assert body.sections[0].heading == "S1"
    assert body.table is None


def test_parse_body_strips_markdown_fences():
    payload = {"title": "Hi", "sections": [{"heading": "S1", "paragraphs": ["p"]}]}
    fenced = f"```json\n{json.dumps(payload)}\n```"
    body = _parse_body(fenced, _make_spec())
    assert body.title == "Hi"


def test_parse_body_raises_on_invalid_json():
    with pytest.raises(RuntimeError, match="non-JSON"):
        _parse_body("not json at all", _make_spec())


def test_parse_body_raises_on_schema_mismatch():
    with pytest.raises(RuntimeError, match="DocBody"):
        _parse_body(json.dumps({"title": "Hi"}), _make_spec())  # missing sections


# ─── Format writers ──────────────────────────────────────────────────────────
def _make_body(with_table: bool = False) -> DocBody:
    return DocBody(
        title="Sample Doc",
        sections=[
            DocSection(heading="Overview", paragraphs=["Para A.", "Para B."]),
            DocSection(heading="Details", paragraphs=["Para C."]),
        ],
        table=(
            DocTable(headers=["col1", "col2"], rows=[["a", "b"], ["c", "d"]])
            if with_table
            else None
        ),
    )


def test_write_md_includes_title_and_sections(tmp_path):
    spec = _make_spec()
    path = tmp_path / "out.md"
    _write_md(path, _make_body(), spec)
    text = path.read_text()
    assert "# Sample Doc" in text
    assert "## Overview" in text
    assert "Para A." in text


def test_write_md_renders_table(tmp_path):
    path = tmp_path / "with_table.md"
    _write_md(path, _make_body(with_table=True), _make_spec(has_table=True))
    text = path.read_text()
    assert "| col1 | col2 |" in text
    assert "| a | b |" in text


def test_write_md_includes_date_stamp(tmp_path):
    spec = _make_spec(date_stamp="2025-04-12 14:00 WIB")
    path = tmp_path / "d.md"
    _write_md(path, _make_body(), spec)
    assert "_2025-04-12 14:00 WIB_" in path.read_text()


def test_write_docx_produces_valid_file(tmp_path):
    path = tmp_path / "out.docx"
    _write_docx(path, _make_body(with_table=True), _make_spec(has_table=True))
    assert path.exists()
    assert path.stat().st_size > 0
    # Verify it parses back as a docx
    from docx import Document
    doc = Document(str(path))
    # Title + 2 section headings = 3 headings
    assert any("Sample Doc" in p.text for p in doc.paragraphs)
    assert len(doc.tables) == 1


def test_write_pdf_produces_valid_file(tmp_path):
    path = tmp_path / "out.pdf"
    _write_pdf(path, _make_body(with_table=True), _make_spec(has_table=True))
    assert path.exists()
    assert path.read_bytes().startswith(b"%PDF-")


# ─── End-to-end with mocked LLM ──────────────────────────────────────────────
@pytest.fixture
def mock_llm(monkeypatch):
    """Replace klerk.llm.router.complete with a canned DocBody response."""
    payload = {
        "title": "Mocked Title",
        "sections": [
            {"heading": "Background", "paragraphs": ["Mocked paragraph one."]},
            {"heading": "Decision", "paragraphs": ["Mocked decision text."]},
        ],
        "table": None,
    }

    @dataclass
    class _Choice:
        message: object

    @dataclass
    class _Msg:
        content: str

    fake_response = MagicMock()
    fake_response.choices = [_Choice(message=_Msg(content=json.dumps(payload)))]

    def fake_complete(**_kwargs):
        return fake_response

    monkeypatch.setattr("klerk.llm.router.complete", fake_complete)
    yield


def test_generate_corpus_writes_all_specs(mock_llm, tmp_path):
    rows = generate_corpus(tmp_path, corpus=CORPUS[:3])
    assert len(rows) == 3
    for r in rows:
        assert r["status"] in {"generated", "cached"}
        assert Path(r["path"]).exists()


def test_generate_corpus_skips_existing(mock_llm, tmp_path):
    sub_corpus = CORPUS[:2]
    rows_first = generate_corpus(tmp_path, corpus=sub_corpus)
    assert all(r["status"] == "generated" for r in rows_first)

    rows_second = generate_corpus(tmp_path, corpus=sub_corpus)
    assert all(r["status"] == "cached" for r in rows_second)


def test_generate_corpus_force_regenerates(mock_llm, tmp_path):
    sub_corpus = CORPUS[:1]
    generate_corpus(tmp_path, corpus=sub_corpus)
    rows = generate_corpus(tmp_path, corpus=sub_corpus, skip_existing=False)
    assert rows[0]["status"] == "generated"


def test_generate_corpus_writes_manifest(mock_llm, tmp_path):
    sub_corpus = CORPUS[:2]
    generate_corpus(tmp_path, corpus=sub_corpus)
    manifest_path = tmp_path / "manifest.json"
    assert manifest_path.exists()
    manifest = json.loads(manifest_path.read_text())
    assert manifest["n_specs"] == 2
    assert manifest["n_generated"] == 2
    assert manifest["n_failed"] == 0
    assert len(manifest["docs"]) == 2


def test_generate_corpus_records_failures(monkeypatch, tmp_path):
    def boom(**_kwargs):
        raise RuntimeError("LLM down")

    monkeypatch.setattr("klerk.llm.router.complete", boom)
    rows = generate_corpus(tmp_path, corpus=CORPUS[:1])
    assert rows[0]["status"] == "failed"
    assert "LLM down" in rows[0]["error"]

    manifest = json.loads((tmp_path / "manifest.json").read_text())
    assert manifest["n_failed"] == 1
    assert manifest["errors"]
